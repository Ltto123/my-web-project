"""
生成「从本地代码到公网部署 全流程实战教学」Word 文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    if level == 1:
        hs.font.size = Pt(22)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    elif level == 3:
        hs.font.size = Pt(13)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)

# 代码块样式
cs = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
cs.font.name = 'Consolas'
cs.font.size = Pt(9)
cs.paragraph_format.space_before = Pt(4)
cs.paragraph_format.space_after = Pt(4)
cs.paragraph_format.left_indent = Cm(0.8)

def code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['CodeBlock']
    for i, line in enumerate(text.strip().split('\n')):
        if i > 0:
            p.add_run('\n')
        p.add_run(line)

def note(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.8)
    return p

# ===== 封面 =====
for _ in range(5):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('从本地代码到公网部署\n全流程实战教学')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

doc.add_paragraph()

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('FastAPI + Docker + 腾讯云 · 一站式部署指南')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run(f'实战项目：个人博客系统（FastAPI + SQLite + 原生前端）\n生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

doc.add_page_break()

# ===== 目录页 =====
doc.add_heading('目录', level=1)
toc = [
    '第一章  为什么要容器化？',
    '第二章  项目结构分析',
    '第三章  编写四个容器化文件',
    '  3.1  requirements.txt —— 依赖清单',
    '  3.2  Dockerfile —— 镜像蓝图',
    '  3.3  .dockerignore —— 瘦身清单',
    '  3.4  docker-compose.yml —— 一键启动',
    '第四章  本地构建与验证',
    '第五章  购买与配置云服务器',
    '  5.1  选哪家、选什么配置',
    '  5.2  SSH 连接服务器',
    '  5.3  开放安全组端口',
    '第六章  在服务器上部署',
    '  6.1  安装 Docker',
    '  6.2  上传项目文件',
    '  6.3  构建镜像并启动容器',
    '  6.4  验证公网访问',
    '第七章  后续更新代码的流程',
    '第八章  常见问题排查',
    '附录A  完整文件参考',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

doc.add_page_break()

# ========================================
# 第一章
# ========================================
doc.add_heading('第一章  为什么要容器化？', level=1)

doc.add_paragraph(
    '想象一个场景：你在自己电脑上写了一个博客，Python 3.12 + FastAPI + SQLite，跑得飞起。'
    '现在你想把它部署到云服务器上让别人也能访问。'
)

doc.add_heading('不用 Docker 的痛苦', level=2)
doc.add_paragraph('1. 登录服务器，发现系统是 Ubuntu，你本地是 Windows，命令不一样')
doc.add_paragraph('2. 手动装 Python 3.12 → 版本不对 → 重装 → 心态崩了')
doc.add_paragraph('3. 手动 pip install 一堆依赖 → 版本冲突 → 一个下午过去了')
doc.add_paragraph('4. 终于跑起来了，换服务器又要重来一遍')
doc.add_paragraph('5. 团队其他人用 macOS，环境跟你不一样，各种"我电脑上能跑啊"')

doc.add_heading('用了 Docker 之后', level=2)
doc.add_paragraph('Docker 就像一个"集装箱"：把你的代码 + Python 环境 + 所有依赖全部打包在一起。')
doc.add_paragraph('  ✅ 一次打包，到处运行（你电脑、服务器、任何人电脑）')
doc.add_paragraph('  ✅ 环境完全一致，没有"我电脑上能跑啊"的问题')
doc.add_paragraph('  ✅ 换服务器只需 30 秒：装 Docker → 传文件 → docker compose up -d')
doc.add_paragraph('  ✅ 方便回滚：新版本有问题，一条命令回到旧版本')
doc.add_paragraph('  ✅ 部署到公司服务器、阿里云、腾讯云、AWS，流程完全一样')

doc.add_heading('核心概念速览', level=2)
doc.add_paragraph('  📦 镜像（Image）：一个只读模板，包含运行环境 + 代码。类比：Windows 安装盘 ISO')
doc.add_paragraph('  🚀 容器（Container）：从镜像启动的运行实例。类比：你用 ISO 装好的、正在跑的系统')
doc.add_paragraph('  📄 Dockerfile：描述"如何一步一步构建镜像"的指令文件。类比：安装说明书')
doc.add_paragraph('  ⚙️  docker-compose.yml：描述"启动容器时用哪些配置"的文件。类比：快捷方式属性')

doc.add_page_break()

# ========================================
# 第二章
# ========================================
doc.add_heading('第二章  项目结构分析', level=1)

doc.add_paragraph('以我们这次部署的博客项目为例，它的目录结构如下：')

code('''my-project/
├── backend/              ← FastAPI 后端代码
│   └── main.py           ← 入口：所有 API 路由 + 静态文件服务
├── frotend/              ← 前端（原生 HTML/CSS/JS）
├── uploads/              ← 用户上传的图片和文件
├── blog.db               ← SQLite 数据库文件（所有博客数据）
├── .env                  ← 环境变量（BLOG_OWNER_USERNAME=博主名）
├── .venv/                ← 本地虚拟环境（⚠️ 不进镜像！容器里自己装）
├── .git/                 ← Git 仓库（⚠️ 不进镜像！）
└── openspec/             ← 内部开发文档（⚠️ 不进镜像！）''')

doc.add_heading('容器化之前，回答三个问题', level=2)
doc.add_paragraph('这是通用方法论，适用于任何项目。')

p = doc.add_paragraph()
r = p.add_run('问题 1：这个项目依赖哪些第三方包？')
r.bold = True
doc.add_paragraph('  → fastapi, uvicorn, sqlalchemy, bcrypt, python-dotenv, python-multipart, pydantic')
doc.add_paragraph('  → 要做什么：把这些写进 requirements.txt')

p = doc.add_paragraph()
r = p.add_run('问题 2：哪些文件需要打包进镜像？')
r.bold = True
doc.add_paragraph('  → backend/（后端代码）+ frotend/（前端文件）+ requirements.txt')
doc.add_paragraph('  → 要做什么：写 Dockerfile 指定来源，写 .dockerignore 排除垃圾')

p = doc.add_paragraph()
r = p.add_run('问题 3：哪些数据需要在容器删除后还保留？')
r.bold = True
doc.add_paragraph('  → blog.db（数据库）和 uploads/（上传文件）')
doc.add_paragraph('  → 要做什么：在 docker-compose.yml 中用 volume 挂载到宿主机')

doc.add_page_break()

# ========================================
# 第三章
# ========================================
doc.add_heading('第三章  编写四个容器化文件', level=1)

doc.add_paragraph('你需要创建四个文件。一句话总结各自的作用：')
doc.add_paragraph('  • requirements.txt     — "我的项目需要这些 Python 包"')
doc.add_paragraph('  • Dockerfile           — "如何一步一步构建出我的镜像"')
doc.add_paragraph('  • .dockerignore        — "这些文件和目录别打包进镜像"')
doc.add_paragraph('  • docker-compose.yml   — "启动容器时用这些配置"')

# ----- 3.1 -----
doc.add_heading('3.1  requirements.txt —— 依赖清单', level=2)

doc.add_paragraph('先去查一下项目到底用了哪些第三方包（在你的本地虚拟环境中）：')
code('pip freeze | grep -iE "fastapi|uvicorn|sqlalchemy|bcrypt|dotenv|multipart"')

doc.add_paragraph('输出类似：')
code('''fastapi==0.136.3
uvicorn==0.48.0
sqlalchemy==2.0.50
bcrypt==5.0.0
python-dotenv==1.2.2
python-multipart==0.0.32
pydantic==2.13.4''')

doc.add_paragraph('把上面这些写进 requirements.txt，最终内容：')
code('''fastapi==0.136.3
uvicorn[standard]==0.48.0
sqlalchemy==2.0.50
bcrypt==5.0.0
python-dotenv==1.2.2
python-multipart==0.0.32
pydantic==2.13.4''')

doc.add_paragraph('📌 关键要点：')
doc.add_paragraph('  • 必须锁定版本号（==x.x.x），否则在另一个环境安装时可能拿到不兼容的新版')
doc.add_paragraph('  • uvicorn 后面加 [standard]：安装了额外的 websockets、httptools 等高性能组件')
doc.add_paragraph('  • Python 标准库不用写（datetime, json, os, uuid, pathlib 是 Python 自带的）')
doc.add_paragraph('  • pydantic 是 fastapi 的依赖，虽然你的代码没直接 import，但 schemas.py 用它定义数据模型')

# ----- 3.2 -----
doc.add_heading('3.2  Dockerfile —— 镜像蓝图', level=2)

doc.add_paragraph('这是四个文件中最核心的一个。它告诉 Docker 如何从零构建你的应用。完整内容：')

code('''# ① 基础镜像：选一个已经装好 Python 3.12 的迷你 Linux
FROM python:3.12-slim

# ② 工作目录：容器内所有后续操作的根路径
WORKDIR /app

# ③ 先单独拷贝依赖文件（利用 Docker 的层缓存机制）
COPY requirements.txt .

# ④ 安装依赖
RUN pip install --no-cache-dir -r requirements.txt

# ⑤ 拷贝项目源代码
COPY backend/ ./backend/
COPY frotend/ ./frotend/

# ⑥ 声明容器运行时监听的端口（文档作用，不实际开放端口）
EXPOSE 8000

# ⑦ 容器启动时执行的命令
CMD ["uvicorn", "backend.main:app", "--host", "0.0.0.0", "--port", "8000"]''')

doc.add_heading('逐行深度讲解', level=3)

doc.add_paragraph('① FROM python:3.12-slim')
doc.add_paragraph('  选基础镜像就像装系统时选 ISO 镜像。python:3.12-slim 是一个已经装好 Python 3.12 的迷你 Linux 系统。')
doc.add_paragraph('  "slim" 表示精简版，体积约 150MB。只包含运行 Python 所需的最小依赖，够用就好。')

doc.add_paragraph('② WORKDIR /app')
doc.add_paragraph('  相当于在容器里执行了 cd /app。之后的 COPY、RUN、CMD 都在这个目录下执行。')
doc.add_paragraph('  如果 /app 目录不存在，Docker 会自动创建。')

doc.add_paragraph('③ COPY requirements.txt .')
doc.add_paragraph('  把宿主机上的 requirements.txt 复制到容器当前目录（/app/）。')
doc.add_paragraph('  📌 为什么要单独 COPY 这一个文件，而不是跟代码一起 COPY？')
doc.add_paragraph('     Docker 构建镜像时是分层（Layer）的。每一层有缓存——如果这层的输入没变化，')
doc.add_paragraph('     Docker 直接复用上次的缓存，跳过执行。')
doc.add_paragraph('     如果你把 COPY backend/ 写在前面，哪怕只改了一行 Python 代码，pip install 层')
doc.add_paragraph('     也会被迫重跑（因为缓存失效了），每次构建多花 30 秒。')
doc.add_paragraph('     而把 COPY requirements.txt 单独放前面 → 只有依赖变了才重装，改代码不重装。')

doc.add_paragraph('④ RUN pip install --no-cache-dir -r requirements.txt')
doc.add_paragraph('  在容器里执行 pip install。--no-cache-dir 表示不保留 pip 下载缓存，让最终镜像更小。')
doc.add_paragraph('  ⚠️ 国内服务器连接 pypi.org 极慢！一定要换镜像源：')
code('RUN pip install --no-cache-dir -i https://pypi.tuna.tsinghua.edu.cn/simple -r requirements.txt')
doc.add_paragraph('  其它可用源：mirrors.aliyun.com/pypi/simple；pypi.douban.com/simple')

doc.add_paragraph('⑤ COPY backend/ ./backend/ 和 COPY frotend/ ./frotend/')
doc.add_paragraph('  把你本地的代码目录复制进容器相应的位置。注意 .venv/、.git/ 等已经通过 .dockerignore 排除了，')
doc.add_paragraph('  Docker 不会拷贝这些目录。')
doc.add_paragraph('  📌 语法：COPY <宿主机路径> <容器路径>。宿主机路径相对于 Dockerfile 所在的目录。')

doc.add_paragraph('⑥ EXPOSE 8000')
doc.add_paragraph(
    '  声明"这个容器运行时会在 8000 端口上监听"。'
    '注意：这只是一个文档声明，并不会真的打开端口！'
)
doc.add_paragraph('  真正让端口对外可访问的，是 docker run 时的 -p 参数，或者 docker-compose.yml 里的 ports 配置。')

doc.add_paragraph('⑦ CMD ["uvicorn", "backend.main:app", "--host", "0.0.0.0", "--port", "8000"]')
doc.add_paragraph('  容器启动时自动执行的命令。拆解如下：')
doc.add_paragraph('    uvicorn                — ASGI 服务器，运行 FastAPI 应用')
doc.add_paragraph('    backend.main:app       — "backend/main.py 文件里的 app 变量"')
doc.add_paragraph('    --host 0.0.0.0         — 监听所有网络接口（默认 127.0.0.1 只接受容器内部请求）')
doc.add_paragraph('    --port 8000            — 监听 8000 端口')
doc.add_paragraph('  📌 为什么用 JSON 数组格式（["cmd", "arg1", "arg2"]）而不是字符串？')
doc.add_paragraph('     JSON 数组格式是 exec 模式，Docker 直接启动进程，信号能正确传递。')
doc.add_paragraph('     字符串格式会套一层 shell，多一个进程，且 Ctrl+C 信号传不进去。始终用数组格式。')

# ----- 3.3 -----
doc.add_heading('3.3  .dockerignore —— 瘦身清单', level=2)

doc.add_paragraph('docker build 时，Docker 会把整个项目目录打包发给 Docker 引擎（叫 build context）。')
doc.add_paragraph('.dockerignore 相当于 .gitignore 的 Docker 版——列在里面的文件和目录不会发送。')
doc.add_paragraph('好处：构建更快（减少传输），镜像更小（不会意外打入垃圾文件）。')

code('''.venv/           # 虚拟环境（几百 MB！容器里自己 pip install）
.git/            # Git 仓库历史（不进镜像）
blog.db          # SQLite 数据库文件（通过 volume 挂载）
uploads/         # 用户上传文件（通过 volume 挂载）
.env             # 本地环境变量（通过 docker-compose 传入）
.vscode/         # VS Code 配置
.claude/         # Claude AI 配置
openspec/        # 内部开发文档
__pycache__/     # Python 字节码缓存
*.pyc            # 编译后的 Python 文件
*.pyo            # 优化编译文件''')

# ----- 3.4 -----
doc.add_heading('3.4  docker-compose.yml —— 一键启动', level=2)

doc.add_paragraph('如果每次都用 docker run 命令，你要敲：')
code('''docker run -d \\
  -p 8000:8000 \\
  -v /home/ubuntu/my-blog/blog.db:/app/blog.db \\
  -v /home/ubuntu/my-blog/uploads:/app/uploads \\
  -e BLOG_OWNER_USERNAME=Ltto123 \\
  --restart unless-stopped \\
  --name my-blog-container \\
  my-blog''')
doc.add_paragraph('又长又容易出错。docker-compose.yml 把所有配置写成文件，以后只需一句 docker compose up -d。')

code('''services:
  blog:
    build: .                           # 从当前目录的 Dockerfile 构建镜像
    ports:
      - "8000:8000"                    # 端口映射：宿主机端口:容器端口
    volumes:
      - ./blog.db:/app/blog.db         # SQLite 数据持久化：宿主机路径:容器路径
      - ./uploads:/app/uploads         # 上传文件持久化
    environment:
      - BLOG_OWNER_USERNAME=Ltto123    # 环境变量
    restart: unless-stopped            # 崩溃后自动重启''')

doc.add_paragraph('📌 ports 和 volumes 的语法都是 宿主机:容器。左边是物理机/服务器，右边是容器内部。')
doc.add_paragraph('📌 为什么 blog.db 和 uploads 要用 volumes 挂载而不是 COPY 进镜像？')
doc.add_paragraph('   容器是无状态的。删除容器 → 容器里所有数据灰飞烟灭。')
doc.add_paragraph('   volumes 把数据存在宿主机硬盘上，容器删了重建，数据毫发无伤。')

doc.add_page_break()

# ========================================
# 第四章
# ========================================
doc.add_heading('第四章  本地构建与验证', level=1)

doc.add_paragraph('四个文件写好后，不要急着往服务器上搬。先在本地验证能跑通。')

doc.add_heading('第 1 步：构建镜像', level=2)
code('docker compose build')
doc.add_paragraph('或者纯 Docker 命令（如果没有 compose）：')
code('docker build -t my-blog .')
doc.add_paragraph('• -t my-blog：给镜像起名叫 my-blog')
doc.add_paragraph('• .：用当前目录作为 build context（Docker 会找当前目录下的 Dockerfile）')

doc.add_heading('第 2 步：启动容器', level=2)
code('docker compose up -d')
doc.add_paragraph('• up：启动服务  •  -d：后台运行（detach，不占终端）')
doc.add_paragraph('第一次运行时会自动构建镜像（等同先 build 再 up）。加了 --build 参数则强制重建：')
code('docker compose up -d --build')

doc.add_heading('第 3 步：检查状态', level=2)
code('''docker compose ps          # 查看容器是否在运行
docker compose logs -f     # 实时查看日志（Ctrl+C 退出）''')

doc.add_heading('第 4 步：验证功能', level=2)
code('curl http://localhost:8000/api/v1/health')
doc.add_paragraph('返回 {"code":0,"msg":"success","data":"healthy"} → 成功！浏览器打开 http://localhost:8000 看博客。')

doc.add_heading('第 5 步：停止和清理', level=2)
code('''docker compose down          # 停止并删除容器（数据在 volume 中不受影响）
docker compose down -v       # 连 volume 一起删除（⚠️ 数据库和上传文件全没！）''')

doc.add_page_break()

# ========================================
# 第五章
# ========================================
doc.add_heading('第五章  购买与配置云服务器', level=1)

doc.add_heading('5.1  选哪家、选什么配置', level=2)

doc.add_paragraph('主流的云服务商：')
doc.add_paragraph('  国内：阿里云、腾讯云 — 新用户免费试用 1-3 个月，国内访问飞快')
doc.add_paragraph('  国外：Hetzner、Vultr、DigitalOcean — 便宜但国内访问慢')
doc.add_paragraph('  免运维：Railway、Fly.io — 一键部署，自动给 URL，但国内可能访问慢')

doc.add_paragraph('选配置的最小要求（够跑一个博客）：')
doc.add_paragraph('  • 操作系统：Ubuntu 22.04 LTS（64 位）')
doc.add_paragraph('  • CPU/内存：2 核 2GB（1 核 1GB 也能跑，但不宽裕）')
doc.add_paragraph('  • 硬盘：20GB 系统盘足够')
doc.add_paragraph('  • 地域：选离你最近的（降低网络延迟）')

doc.add_paragraph('买完后你手里有这些信息：')
doc.add_paragraph('  • 公网 IP 地址（如 111.229.173.39）← 这是关键')
doc.add_paragraph('  • 默认用户名（腾讯云 ubuntu，阿里云 root）')
doc.add_paragraph('  • 密码（去控制台"重置密码"自己设一个）')

doc.add_heading('5.2  SSH 连接服务器', level=2)

doc.add_paragraph('SSH（Secure Shell）是远程登录 Linux 服务器的标准协议。')

doc.add_paragraph('在终端输入：')
code('ssh 用户名@公网IP')
doc.add_paragraph('例：ssh ubuntu@111.229.173.39')

doc.add_paragraph('第一次连接会提示：')
code('The authenticity of host "111.229.173.39" cannot be established.')
doc.add_paragraph('这是正常的，输入 yes 回车。然后输入密码。')
doc.add_paragraph('📌 输入密码时屏幕上不会显示任何字符（连 *** 都不显示），这是安全设计。输完回车就行。')

doc.add_paragraph('连接成功后，终端提示符会变成 用户名@主机名:~$，恭喜，你现在在服务器里面了！')

doc.add_heading('常见 SSH 错误', level=2)
doc.add_paragraph('  ❌ Permission denied → 密码错了，去控制台重置')
doc.add_paragraph('  ❌ Connection timed out → IP 错了，或者安全组没放行 22 端口')
doc.add_paragraph('  ❌ Host key verification failed → 服务器重装过系统，编辑 ~/.ssh/known_hosts 删掉对应行')

doc.add_heading('5.3  开放安全组端口', level=2)

doc.add_paragraph('⚠️ 这是新人最容易漏掉的一步！漏了这一步，公网永远访问不到你的博客。')

doc.add_paragraph('云服务器有两层防火墙：')
doc.add_paragraph('  第一层：服务器自身的 ufw 防火墙 → 默认关闭，不用管')
doc.add_paragraph('  第二层：云厂商的"安全组" → 在网页控制台配置，必须手动添加规则！')

doc.add_paragraph('腾讯云操作步骤：')
doc.add_paragraph('  1. 登录腾讯云控制台 → 云服务器 → 实例')
doc.add_paragraph('  2. 点击你的服务器 → 上方标签选"安全组"')
doc.add_paragraph('  3. 点击"添加规则"，填写：')
code('''来源：    0.0.0.0/0        （允许所有 IP 地址访问）
协议端口： TCP:8000          （你的博客监听 8000 端口）
策略：     允许''')
doc.add_paragraph('  4. 点击保存，等待 10 秒生效')

doc.add_paragraph('阿里云操作类似，叫"安全组规则"。')
doc.add_paragraph('必须开放两个端口：22（SSH，通常默认已开）和 8000（你的博客）。')

doc.add_page_break()

# ========================================
# 第六章
# ========================================
doc.add_heading('第六章  在服务器上部署', level=1)

doc.add_paragraph('这是全流程中操作最密集的环节，严格按照以下顺序操作。')

doc.add_heading('6.1  安装 Docker', level=2)

doc.add_paragraph('SSH 登录服务器后，依次执行：')
code('''# 1. 更新包列表
sudo apt update

# 2. 安装 Docker 引擎和 Docker Compose 插件
sudo apt install -y docker.io docker-compose-v2

# 3. 启动 Docker 服务并设置开机自动启动
sudo systemctl enable docker --now

# 4. 把当前用户加入 docker 组（之后不用每次 sudo）
sudo usermod -aG docker ubuntu

# 5. ⚠️ 退出 SSH 重新登录，组权限才能生效
exit
ssh ubuntu@你的IP

# 6. 验证安装
docker --version        # 应输出 Docker version xx.xx.x
docker compose version  # 应输出 Docker Compose version vx.xx.x''')

doc.add_paragraph('📌 如果 docker-compose-v2 安装报冲突（常见于腾讯云），用备用方案：')
code('''sudo mkdir -p /usr/local/lib/docker/cli-plugins
sudo curl -SL "https://github.com/docker/compose/releases/latest/download/docker-compose-linux-x86_64" -o /usr/local/lib/docker/cli-plugins/docker-compose
sudo chmod +x /usr/local/lib/docker/cli-plugins/docker-compose
docker compose version   # 验证''')

doc.add_heading('6.2  上传项目文件', level=2)

doc.add_paragraph('在本地电脑上，把项目打包（排除不需要进服务器的垃圾）：')
code('''cd 你的项目目录
tar -czf project.tar.gz \\
  --exclude=".venv" --exclude=".git" --exclude=".claude" \\
  --exclude="blog.db" --exclude="uploads" --exclude="openspec" \\
  --exclude="__pycache__" --exclude="*.pyc" \\
  Dockerfile docker-compose.yml requirements.txt .dockerignore \\
  backend/ frotend/''')

doc.add_paragraph('📌 tar 命令解说：')
doc.add_paragraph('  -c：创建归档  -z：gzip 压缩  -f：指定文件名')
doc.add_paragraph('  --exclude="xxx"：排除指定文件/目录')

doc.add_paragraph('然后用 scp 命令传到服务器：')
code('scp project.tar.gz ubuntu@你的IP:~/')
doc.add_paragraph('📌 scp 语法：scp <本地文件路径> <用户名@服务器IP:服务器目标路径>')
doc.add_paragraph('意思是：把本地的 project.tar.gz 复制到服务器的 /home/ubuntu/ 目录下。')

doc.add_paragraph('传到服务器后，SSH 进去解压：')
code('''ssh ubuntu@你的IP
mkdir -p ~/my-blog
cd ~/my-blog
tar -xzf ~/project.tar.gz
ls -la   # 检查文件都在：Dockerfile, backend/, frotend/...''')
doc.add_paragraph('📌 tar -xzf：x=解压，z=gzip 格式，f=指定文件')

doc.add_heading('6.3  构建镜像并启动容器', level=2)

doc.add_paragraph('⚡ 关键操作：国内服务器 pip 连 pypi.org 很慢。先修改服务器上的 Dockerfile：')
code('''# 把这一行
RUN pip install --no-cache-dir -r requirements.txt
# 改成这一行（加 -i 指定清华源）
RUN pip install --no-cache-dir -i https://pypi.tuna.tsinghua.edu.cn/simple -r requirements.txt''')
doc.add_paragraph('改完后，执行构建和启动：')
code('''cd ~/my-blog

# 创建初始文件（否则 volume 挂载时会把 blog.db 当成目录）
touch blog.db
mkdir -p uploads

# 构建镜像
sudo docker build -t my-blog .

# 启动容器
sudo docker run -d \\
  --name my-blog-container \\
  -p 8000:8000 \\
  -v /home/ubuntu/my-blog/blog.db:/app/blog.db \\
  -v /home/ubuntu/my-blog/uploads:/app/uploads \\
  -e BLOG_OWNER_USERNAME=Ltto123 \\
  --restart unless-stopped \\
  my-blog''')

doc.add_paragraph('逐参数讲解：')
doc.add_paragraph('  -d          后台运行（如果不加，终端会一直被占用，关终端容器就停）')
doc.add_paragraph('  --name      给容器起个名字，方便后续操作（docker stop my-blog-container）')
doc.add_paragraph('  -p 8000:8000  端口映射。左：服务器端口（外部访问这个），右：容器内端口（uvicorn 监听的）')
doc.add_paragraph('  -v /path:/app/blog.db  卷挂载。左：服务器上的真实文件，右：容器内看到的文件路径')
doc.add_paragraph('  -e KEY=VALUE  设置环境变量，这里指定了博主用户名')
doc.add_paragraph('  --restart unless-stopped  容器崩溃或服务器重启后，Docker 自动把容器拉起来')
doc.add_paragraph('  my-blog  使用哪个镜像（刚才 build 出来的）')

doc.add_heading('常用管理命令', level=2)
code('''docker ps                    # 查看正在运行的容器
docker ps -a                 # 查看所有容器（包括已停止的）
docker logs my-blog-container   # 查看容器日志（排查问题必用）
docker restart my-blog-container  # 重启容器
docker rm -f my-blog-container   # 强制删除容器
docker images                # 查看本地镜像列表''')

doc.add_heading('6.4  验证公网访问', level=2)
code('''# 在服务器内部测试
curl http://localhost:8000/api/v1/health

# 在你本地电脑上测试（用服务器的公网 IP）
curl http://111.229.173.39:8000/api/v1/health''')
doc.add_paragraph('如果返回 {"code":0,"msg":"success","data":"healthy"}：恭喜，部署成功！🎉')
doc.add_paragraph('浏览器打开 http://你的公网IP:8000，博客面向全世界开放了。')

doc.add_page_break()

# ========================================
# 第七章
# ========================================
doc.add_heading('第七章  后续更新代码的流程', level=1)

doc.add_paragraph('写了一个新功能，改了几个 bug，怎么更新到服务器上？')

doc.add_heading('标准更新流程（5 步，约 30 秒）', level=2)
code('''# ===== 在你本地电脑上 =====

# 1. 重新打包代码（⚠️ 不要打包 .venv、blog.db、uploads）
tar -czf project.tar.gz \\
  --exclude=".venv" --exclude=".git" --exclude=".claude" \\
  --exclude="blog.db" --exclude="uploads" --exclude="openspec" \\
  --exclude="__pycache__" --exclude="*.pyc" \\
  Dockerfile docker-compose.yml requirements.txt .dockerignore \\
  backend/ frotend/

# 2. 上传到服务器
scp project.tar.gz ubuntu@你的IP:~/my-blog/

# ===== SSH 进入服务器 =====
ssh ubuntu@你的IP
cd ~/my-blog

# 3. 解压，覆盖旧代码
tar -xzf project.tar.gz

# 4. 重新构建镜像
sudo docker build -t my-blog .

# 5. 用新镜像重启容器
sudo docker rm -f my-blog-container
sudo docker run -d --name my-blog-container \\
  -p 8000:8000 \\
  -v /home/ubuntu/my-blog/blog.db:/app/blog.db \\
  -v /home/ubuntu/my-blog/uploads:/app/uploads \\
  -e BLOG_OWNER_USERNAME=Ltto123 \\
  --restart unless-stopped \\
  my-blog''')

doc.add_paragraph('📌 blog.db 和 uploads 用了 volume 挂载，所以删容器重开不会丢数据。')
doc.add_paragraph('📌 整个过程最多中断服务 3-5 秒（docker rm 到新容器启动）。')

doc.add_page_break()

# ========================================
# 第八章
# ========================================
doc.add_heading('第八章  常见问题排查', level=1)

problems = [
    ('容器启动后立刻退出',
     '用 docker ps -a 查看状态。如果 STATUS 显示 "Exited"，用 docker logs 容器名 查看错误日志。\n'
     '常见原因：Python import 报错（路径不对）、端口被占用、数据库文件权限不够。'),
    ('端口被占用',
     'Error: "port is already allocated"\n'
     '用 sudo lsof -i :8000 看谁在占用 → sudo kill PID 干掉它。\n'
     '或者 sudo docker rm -f 旧容器名，强制删除。'),
    ('pip install 超时或极慢',
     '国内服务器连 pypi.org 基本不通或极慢。\n'
     '解决：在 Dockerfile 里的 pip install 加入 -i https://pypi.tuna.tsinghua.edu.cn/simple\n'
     '其它可用源：mirrors.aliyun.com/pypi/simple、pypi.douban.com/simple'),
    ('docker compose up 报错找不到文件',
     '确认在 Dockerfile 所在的目录执行命令。\n'
     '确认 docker-compose.yml 里的 build: . 指向正确的目录。'),
    ('公网 IP 能 ping 通但浏览器打不开',
     '99% 是安全组没放行 8000 端口。去云厂商控制台检查。\n'
     '还有 1% 可能：服务器 ufw 拦了 → sudo ufw allow 8000/tcp。'),
    ('blog.db 数据丢了',
     '检查启动容器时有没有加 -v 参数挂载。没挂载的话，容器一删数据全没。\n'
     '验证挂载：docker inspect 容器名 | grep -A 5 Mounts'),
    ('容器内部 localhost:8000 通，外部不通',
     '检查启动命令里有没有 --host 0.0.0.0。默认 uvicorn 只监听 127.0.0.1，不接受外部请求。\n'
     '检查 docker run 有没有 -p 8000:8000 端口映射。'),
    ('Docker 拉取镜像失败',
     '国内拉取 Docker Hub 很可能失败。\n'
     '解决：配置 Docker 镜像加速器。编辑 /etc/docker/daemon.json，加入：'),
]

for title, solution in problems:
    doc.add_heading(title, level=2)
    if 'daemon.json' in solution:
        code('''{
  "registry-mirrors": [
    "https://docker.m.daocloud.io",
    "https://docker.1ms.run"
  ]
}''')
        doc.add_paragraph('然后 sudo systemctl restart docker。')
    doc.add_paragraph(solution)

doc.add_page_break()

# ========================================
# 附录A
# ========================================
doc.add_heading('附录A  完整文件参考', level=1)

doc.add_paragraph('可以直接复制粘贴使用的四个文件完整内容：')

doc.add_heading('requirements.txt', level=2)
code('''fastapi==0.136.3
uvicorn[standard]==0.48.0
sqlalchemy==2.0.50
bcrypt==5.0.0
python-dotenv==1.2.2
python-multipart==0.0.32
pydantic==2.13.4''')

doc.add_heading('Dockerfile', level=2)
code('''FROM python:3.12-slim
WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt
COPY backend/ ./backend/
COPY frotend/ ./frotend/
EXPOSE 8000
CMD ["uvicorn", "backend.main:app", "--host", "0.0.0.0", "--port", "8000"]''')

doc.add_heading('.dockerignore', level=2)
code('''.venv/
.git/
.gitignore
.vscode/
.idea/
.claude/
openspec/
blog.db
uploads/
.env
__pycache__/
*.pyc
*.pyo''')

doc.add_heading('docker-compose.yml', level=2)
code('''services:
  blog:
    build: .
    ports:
      - "8000:8000"
    volumes:
      - ./blog.db:/app/blog.db
      - ./uploads:/app/uploads
    environment:
      - BLOG_OWNER_USERNAME=Ltto123
    restart: unless-stopped''')

# 结尾
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— 全文完 —')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)
r.italic = True

# ===== 保存 =====
output_path = r'C:\Users\dudu1\Desktop\从本地代码到公网部署_全流程实战教学.docx'
doc.save(output_path)
print(f'Done: {output_path}')
